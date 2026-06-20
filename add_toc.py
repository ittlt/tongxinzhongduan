#!/usr/bin/env python3
"""Insert a TOC (Table of Contents) field into the Word document."""

import copy
from docx import Document
from docx.oxml.ns import qn, nsmap
from lxml import etree

doc = Document('DDS_Signal_Generator_结题报告_Zynq.docx')

# Find the paragraph that contains "目录"
toc_heading = None
toc_text_para = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '目录':
        toc_heading = para
        break

if toc_heading is None:
    print("ERROR: '目录' heading not found")
    exit(1)

print(f"Found '目录' at paragraph index {i}")

# Remove the manual TOC text paragraphs that follow the "目录" heading
# We need to find and remove the paragraphs between "目录" and "项目背景与需求"
paragraphs_to_remove = []
found_toc = False
for j, para in enumerate(doc.paragraphs):
    if para.text.strip() == '目录':
        found_toc = True
        continue
    if found_toc:
        if '项目背景与需求' in para.text:
            break
        # Check if this is a TOC entry paragraph
        text = para.text.strip()
        if text.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、')):
            paragraphs_to_remove.append(para)
        elif text.startswith(('1.', '2.', '3.', '4.', '5.')):
            paragraphs_to_remove.append(para)
        elif text == '':
            paragraphs_to_remove.append(para)

print(f"Found {len(paragraphs_to_remove)} TOC paragraphs to remove")

# Remove the manual TOC paragraphs by removing their XML elements
for para in paragraphs_to_remove:
    p = para._element
    p.getparent().remove(p)

print("Removed manual TOC paragraphs")

# Now insert a TOC field after the "目录" heading
# The TOC field in Word XML
TOC_INSTR = r' TOC \o "1-3" \h \z \u '

# Create the TOC field XML
fldChar_begin = etree.SubElement(etree.Element('dummy'), qn('w:fldChar'))
fldChar_begin.set(qn('w:fldCharType'), 'begin')

fldChar_separate = etree.SubElement(etree.Element('dummy'), qn('w:fldChar'))
fldChar_separate.set(qn('w:fldCharType'), 'separate')

fldChar_end = etree.SubElement(etree.Element('dummy'), qn('w:fldChar'))
fldChar_end.set(qn('w:fldCharType'), 'end')

# Get the XML element of the toc_heading paragraph
toc_p = toc_heading._element

# Create run with field begin
run_begin = etree.SubElement(toc_p, qn('w:r'))
rPr = etree.SubElement(run_begin, qn('w:rPr'))
fldChar = etree.SubElement(run_begin, qn('w:fldChar'))
fldChar.set(qn('w:fldCharType'), 'begin')

# Create run with instruction text
run_instr = etree.SubElement(toc_p, qn('w:r'))
rPr2 = etree.SubElement(run_instr, qn('w:rPr'))
instrText = etree.SubElement(run_instr, qn('w:instrText'))
instrText.set(qn('xml:space'), 'preserve')
instrText.text = TOC_INSTR

# Create run with field separate
run_sep = etree.SubElement(toc_p, qn('w:r'))
rPr3 = etree.SubElement(run_sep, qn('w:rPr'))
fldChar2 = etree.SubElement(run_sep, qn('w:fldChar'))
fldChar2.set(qn('w:fldCharType'), 'separate')

# Create placeholder text
run_placeholder = etree.SubElement(toc_p, qn('w:r'))
rPr4 = etree.SubElement(run_placeholder, qn('w:rPr'))
t = etree.SubElement(run_placeholder, qn('w:t'))
t.text = '(请在Word中右键点击目录，选择"更新域"以生成目录)'

# Create run with field end
run_end = etree.SubElement(toc_p, qn('w:r'))
rPr5 = etree.SubElement(run_end, qn('w:rPr'))
fldChar3 = etree.SubElement(run_end, qn('w:fldChar'))
fldChar3.set(qn('w:fldCharType'), 'end')

# Save the document
doc.save('DDS_Signal_Generator_结题报告_Zynq.docx')
print("Document saved with TOC field")
print("Note: Please open in Word and update the TOC field to generate the clickable table of contents")
