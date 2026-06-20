#!/usr/bin/env python3
"""Insert the system block diagram at the correct location in the document."""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import os

doc = Document('DDS_Signal_Generator_结题报告_Zynq.docx')

# Find the "系统设计总框图" paragraph
target_para = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '系统设计总框图':
        target_para = para
        print(f"Found '系统设计总框图' at paragraph index {i}")
        break

if target_para is None:
    print("ERROR: '系统设计总框图' paragraph not found!")
    exit(1)

# Find the next paragraph after "系统设计总框图" which should be the text description
next_para = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('系统设计总框图如下所示'):
        next_para = para
        print(f"Found text description at paragraph index {i}")
        break

# Add the picture after "系统设计总框图" heading
# We need to insert a new paragraph with the image after the target_para
from docx.oxml.ns import nsmap
from docx.shared import Inches

# Get the parent element
body = doc.element.body

# Create a new paragraph element
new_p = etree.SubElement(body, qn('w:p'))

# Find the position of target_para and insert after it
target_elem = target_para._element
target_idx = list(body).index(target_elem)

# Insert the new paragraph after the target paragraph
body.insert(target_idx + 1, new_p)

# Now add the image to this new paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.image.image import Image
from docx.shared import Emu

# Add relationship for the image
image_path = '/mnt/d/FPGAmoudle/--main/--main/system_block_diagram_cn.png'

# Get the image size
img = Image(image_path)
width_emu = int(6 * 914400)  # 6 inches in EMU

# Add the picture
new_para = doc.paragraphs[target_idx + 1] if target_idx + 1 < len(doc.paragraphs) else None

if new_para is not None:
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(6))
    print("Picture inserted successfully!")
else:
    print("ERROR: Could not add picture to new paragraph")

# Save the document
doc.save('DDS_Signal_Generator_结题报告_Zynq.docx')
print("Document saved")
