#!/usr/bin/env python3
"""Fix the TOC in the Word document - remove manual entries, keep only the TOC field."""

from docx import Document
from lxml import etree

doc = Document('DDS_Signal_Generator_结题报告_Zynq.docx')

# Find the "目录" paragraph and remove everything between it and "项目背景与需求"
toc_index = None
body_index = None

for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '目录' or '目录' in para.text and 'TOC' not in para.text:
        # Check if this paragraph has a TOC field
        has_toc_field = False
        for run in para.runs:
            for child in run._element:
                if child.tag.endswith('}fldChar') or child.tag.endswith('}instrText'):
                    has_toc_field = True
                    break
        if not has_toc_field:
            toc_index = i
    if para.text.strip() == '项目背景与需求':
        body_index = i
        break

print(f"TOC heading at index: {toc_index}")
print(f"Body start at index: {body_index}")

if toc_index is not None and body_index is not None:
    # Remove paragraphs between TOC heading and body
    # But we need to be careful - some paragraphs might be part of the TOC field
    paragraphs_to_remove = []
    for i in range(toc_index + 1, body_index):
        para = doc.paragraphs[i]
        text = para.text.strip()
        # Check if this is a manual TOC entry
        if text.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、')):
            paragraphs_to_remove.append(para)
        elif text.startswith(('1.', '2.', '3.', '4.', '5.')):
            paragraphs_to_remove.append(para)
        elif text == '':
            paragraphs_to_remove.append(para)

    print(f"Removing {len(paragraphs_to_remove)} manual TOC paragraphs")
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

# Now let's also make sure the TOC field paragraph is clean
# Find the paragraph that contains "目录" and has the TOC field
for i, para in enumerate(doc.paragraphs):
    if '目录' in para.text:
        # Check if this has the TOC field
        has_toc = False
        for run in para.runs:
            for child in run._element:
                if child.tag.endswith('}instrText') and 'TOC' in (child.text or ''):
                    has_toc = True
                    break
        if has_toc:
            print(f"Found TOC field paragraph at index {i}")
            # Remove any text that's not part of the TOC field
            # The paragraph text should only be the placeholder
            break

doc.save('DDS_Signal_Generator_结题报告_Zynq.docx')
print("Document saved")
