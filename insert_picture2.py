#!/usr/bin/env python3
"""Insert the system block diagram at the correct location in the document."""

from docx import Document
from docx.shared import Inches

doc = Document('DDS_Signal_Generator_结题报告_Zynq.docx')

# Find the "系统设计总框图" paragraph
target_para = None
target_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '系统设计总框图':
        target_para = para
        target_idx = i
        print(f"Found '系统设计总框图' at paragraph index {i}")
        break

if target_para is None:
    print("ERROR: '系统设计总框图' paragraph not found!")
    exit(1)

# Find the text description paragraph
text_para = None
text_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('系统设计总框图如下所示'):
        text_para = para
        text_idx = i
        print(f"Found text description at paragraph index {i}")
        break

# Add the picture after the "系统设计总框图" heading
# We'll add a new paragraph with the image
image_path = '/mnt/d/FPGAmoudle/--main/--main/system_block_diagram_cn.png'

# Add a new paragraph after the target paragraph
new_para = doc.add_paragraph()
new_para.alignment = 1  # Center alignment

# Move the new paragraph to the correct position
# Get the XML elements
target_elem = target_para._element
new_elem = new_para._element
parent = target_elem.getparent()

# Remove the new paragraph from its current position
parent.remove(new_elem)

# Find the index of target_elem and insert after it
target_pos = list(parent).index(target_elem)
parent.insert(target_pos + 1, new_elem)

# Now add the image to this paragraph
run = new_para.add_run()
run.add_picture(image_path, width=Inches(6))

print("Picture inserted successfully!")

# Save the document
doc.save('DDS_Signal_Generator_结题报告_Zynq.docx')
print("Document saved")
